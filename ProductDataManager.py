import streamlit as st
import pandas as pd
from docx import Document
import io
import re
import json
from datetime import datetime
import requests
import pdfplumber


class LLMClient:
    """Универсальный клиент для работы с Ollama или LM Studio"""

    def __init__(self, provider="ollama"):
        self.provider = provider
        if provider == "ollama":
            self.base_url = "http://localhost:11434"
            self.model = "mistral:7b-instruct"
        else:
            self.base_url = "http://localhost:1234/v1"
            self.model = "phi-3-mini"

    def is_available(self):
        try:
            if self.provider == "ollama":
                response = requests.get(f"{self.base_url}/api/tags", timeout=5)
                return response.status_code == 200
            else:
                response = requests.get(f"{self.base_url}/models", timeout=5)
                return response.status_code == 200
        except requests.exceptions.RequestException:
            return False

    def generate_response(self, prompt):
        if self.provider == "ollama":
            return self._ollama_generate(prompt)
        else:
            return self._lm_studio_generate(prompt)

    def _ollama_generate(self, prompt):
        payload = {
            "model": self.model,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0.1, "num_predict": 1000}
        }
        try:
            response = requests.post(f"{self.base_url}/api/generate", json=payload, timeout=120)
            if response.status_code == 200:
                return response.json().get("response", "").strip()
            return None
        except requests.exceptions.RequestException as e:
            st.error(f"Ошибка соединения с Ollama: {e}")
            return None

    def _lm_studio_generate(self, prompt):
        payload = {
            "model": self.model,
            "messages": [
                {"role": "system",
                 "content": "Ты полезный ассистент для подбора товаров. Строго следуй правилам совместимости."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1,
            "max_tokens": 1000
        }
        try:
            response = requests.post(
                f"{self.base_url}/chat/completions",
                json=payload,
                timeout=120,
                headers={"Content-Type": "application/json"}
            )
            if response.status_code == 200:
                result = response.json()
                return result['choices'][0]['message']['content'].strip()
            return None
        except requests.exceptions.RequestException as e:
            st.error(f"Ошибка соединения с LM Studio: {e}")
            return None


class DataLoader:
    def __init__(self):
        self.products_df = None
        self.compatibility_df = None  # Новое: для хранения правил

    def load_products_from_files(self, uploaded_files):
        try:
            all_dfs = []
            for uploaded_file in uploaded_files:
                file_name = uploaded_file.name.lower()

                # СПЕЦИАЛЬНАЯ ОБРАБОТКА compatibility.csv
                if file_name == 'compatibility.csv':
                    try:
                        self.compatibility_df = pd.read_csv(uploaded_file)
                        st.info(f"📜 Загружены правила совместимости: {len(self.compatibility_df)} шт.")
                    except Exception as e:
                        st.error(f"Ошибка чтения {uploaded_file.name}: {e}")
                    continue  # Переходим к следующему файлу, не добавляем в товары

                if file_name.endswith('.csv'):
                    df = self._load_csv(uploaded_file)
                elif file_name.endswith('.pdf'):
                    df = self._load_pdf(uploaded_file)
                elif file_name.endswith('.json'):
                    st.warning(f"JSON файлы обрабатываются в отдельной вкладке. Файл {uploaded_file.name} пропущен.")
                    continue
                else:
                    st.error(f"Неподдерживаемый формат: {uploaded_file.name}")
                    continue

                if df is not None and not df.empty:
                    all_dfs.append(df)

            if all_dfs:
                self.products_df = pd.concat(all_dfs, ignore_index=True)
                self.products_df = self.products_df.drop_duplicates(subset=['Товар'])
                return True
            return False

        except Exception as e:
            st.error(f"Критическая ошибка загрузки файлов: {str(e)}")
            return False

    def _load_csv(self, uploaded_file):
        try:
            df = pd.read_csv(uploaded_file)
            if 'Товар' not in df.columns:
                for col in df.columns:
                    if any(keyword in str(col).lower() for keyword in
                           ['товар', 'product', 'наименование', 'name', 'description']):
                        df = df.rename(columns={col: 'Товар'})
                        break
                else:
                    df.columns = ['Товар'] + list(df.columns[1:])
            return df
        except Exception as e:
            st.error(f"Ошибка чтения CSV {uploaded_file.name}: {e}")
            return None

    def _load_pdf(self, uploaded_file):
        try:
            with pdfplumber.open(uploaded_file) as pdf:
                text_content = [page.extract_text() for page in pdf.pages if page.extract_text()]

            full_text = "\n".join(text_content)
            product_lines = []
            lines = full_text.splitlines()
            current_line = ""

            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if 'руб' in line.lower():
                    if current_line:
                        line = current_line + " " + line
                        current_line = ""
                    clean_line = re.sub(r'\s+', ' ', line.strip())
                    product_lines.append(clean_line)
                else:
                    current_line = line

            if current_line and 'руб' in current_line.lower():
                product_lines.append(re.sub(r'\s+', ' ', current_line.strip()))

            if product_lines:
                return pd.DataFrame(product_lines, columns=['Товар'])
            else:
                st.warning(f"В PDF {uploaded_file.name} не найдено строк с 'руб'")
                return None
        except Exception as e:
            st.error(f"Ошибка чтения PDF {uploaded_file.name}: {e}")
            return None

    def parse_product_info(self, product_string):
        try:
            product_str = str(product_string)
            price_match = re.search(r'(\d+(?:\s?\d+)*\s*руб\.?)', product_str, re.IGNORECASE)
            if price_match:
                price_str = price_match.group(1)
                clean_price = re.sub(r'\D', '', price_str)
                price = int(clean_price) if clean_price else 0
            else:
                price = 0
            return {'name': product_str, 'price': price}
        except Exception:
            return {'name': str(product_string), 'price': 0}


class AISearch:
    def __init__(self, llm_client, data_loader):
        self.llm_client = llm_client
        self.data_loader = data_loader

    def find_products(self, query):
        if not self.llm_client.is_available():
            st.error("LLM сервис не доступен. Убедитесь, что он запущен.")
            return []

        if self.data_loader.products_df is None or self.data_loader.products_df.empty:
            st.error("Нет загруженных данных о товарах")
            return []

        # 🧠 УМНЫЙ ПОИСК (Lightweight RAG): фильтруем товары по ключевым словам
        all_products = [str(p) for p in self.data_loader.products_df['Товар']]
        query_words = [w.lower() for w in query.split() if len(w) > 2]

        relevant_products = []
        for p in all_products:
            p_lower = p.lower()
            # Если хотя бы одно слово из запроса есть в товаре, считаем его релевантным
            if any(word in p_lower for word in query_words):
                relevant_products.append(p)

        # Если нашли слишком много, ограничиваем топ-30. Если мало, добавляем первые товары для контекста
        if len(relevant_products) > 30:
            relevant_products = relevant_products[:30]
        elif len(relevant_products) < 10:
            relevant_products = list(
                dict.fromkeys(relevant_products + all_products[:30]))  # dict.fromkeys убирает дубликаты

        products_text = "\n".join([f"{i + 1}. {p}" for i, p in enumerate(relevant_products)])

        # 🔗 ДОБАВЛЯЕМ ПРАВИЛА СОВМЕСТИМОСТИ В ПРОМПТ
        compatibility_context = ""
        if self.data_loader.compatibility_df is not None:
            matching_rules = []
            for _, row in self.data_loader.compatibility_df.iterrows():
                rule = str(row['rule'])
                if any(word in rule.lower() for word in query_words):
                    matching_rules.append(f"- {rule}")

            if matching_rules:
                compatibility_context = "\n❗ ВАЖНЫЕ ПРАВИЛА СОВМЕСТИМОСТИ (обязательно учитывай их):\n" + "\n".join(
                    matching_rules)

        prompt = f"""
        Запрос клиента: "{query}"
        {compatibility_context}

        Доступные товары (наиболее релевантные из каталога):
        {products_text}

        Выбери от 1 до 5 наиболее подходящих товаров для этого запроса, строго учитывая правила совместимости.
        Верни ТОЛЬКО номера выбранных товаров через запятую в формате: 1, 3, 5
        Не добавляй никакого дополнительного текста, объяснений или скобок.
        """

        response = self.llm_client.generate_response(prompt)

        if response:
            numbers = re.findall(r'\b\d+\b', response)
            selected_indices = []
            for num in numbers:
                idx = int(num) - 1
                if 0 <= idx < len(relevant_products):
                    selected_indices.append(idx)

            if selected_indices:
                matches = []
                seen = set()
                for idx in selected_indices:
                    if idx not in seen:
                        seen.add(idx)
                        product_info = self.data_loader.parse_product_info(relevant_products[idx])
                        matches.append(product_info)
                return matches[:5]

        # Fallback, если LLM вернула мусор
        return self._fallback_search(query, all_products)

    def _fallback_search(self, query, products_list):
        query_words = [w for w in query.lower().split() if len(w) > 2]
        matches = []
        for product_str in products_list:
            product_lower = product_str.lower()
            score = sum(1 for word in query_words if word in product_lower)
            if score > 0:
                product_info = self.data_loader.parse_product_info(product_str)
                matches.append((score, product_info))
        matches.sort(key=lambda x: x[0], reverse=True)
        return [product for score, product in matches[:5]]


class DocumentGenerator:
    def generate_tcp_document(self, order_data, products):
        doc = Document()
        doc.add_heading('Технико-коммерческое предложение', 0)

        order_paragraph = doc.add_paragraph()
        order_paragraph.add_run("Заказ № ").bold = True
        order_paragraph.add_run(f"{order_data['order_id']}")
        order_paragraph.add_run(f" от {datetime.now().strftime('%d.%m.%Y')}")

        doc.add_paragraph("Перечень комплектующих:")

        total_cost = 0
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Наименование товара'
        hdr_cells[1].text = 'Стоимость'

        for product in products:
            row_cells = table.add_row().cells
            row_cells[0].text = product['name']
            row_cells[1].text = f"{product['price']} руб."
            total_cost += product['price']

        doc.add_paragraph("\n")
        price_paragraph = doc.add_paragraph()
        price_paragraph.add_run(f"ИТОГО: {total_cost} руб.").bold = True

        return doc, total_cost


class JSONProcessor:
    def __init__(self, ai_search, doc_generator):
        self.ai_search = ai_search
        self.doc_generator = doc_generator

    def process_json_file(self, json_file):
        try:
            if json_file.size > 5 * 1024 * 1024:
                st.error(f"Файл {json_file.name} слишком большой (макс. 5 MB)")
                return []

            json_file.seek(0)
            json_data = json.load(json_file)

            if isinstance(json_data, list):
                return self._process_multiple_queries(json_data)
            elif isinstance(json_data, dict):
                if 'query' in json_data:
                    return self._process_single_query(json_data)
                elif 'queries' in json_data:
                    return self._process_multiple_queries(json_data['queries'])
                else:
                    st.error(f"Некорректная структура JSON в {json_file.name}: нет 'query' или 'queries'")
                    return []
            else:
                st.error(f"Некорректный формат JSON в {json_file.name}")
                return []
        except json.JSONDecodeError as e:
            st.error(f"Ошибка парсинга JSON в {json_file.name}: {str(e)}")
            return []
        except Exception as e:
            st.error(f"Ошибка обработки {json_file.name}: {str(e)}")
            return []

    def _process_single_query(self, query_data):
        query_text = str(query_data.get('query', ''))
        query_id = str(query_data.get('id', 'unknown'))
        if not query_text:
            return []

        matches = self.ai_search.find_products(query_text)
        if not matches:
            return []

        order_data = {'order_id': f"JSON-{query_id}"}
        doc, total_cost = self.doc_generator.generate_tcp_document(order_data, matches)

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return [{
            'query_id': query_id,
            'query_text': query_text,
            'matches': matches,
            'total_cost': total_cost,
            'document': doc_buffer,
            'filename': f"tcp_json_{query_id}.docx"
        }]

    def _process_multiple_queries(self, queries):
        results = []
        for query_data in queries:
            if isinstance(query_data, dict) and 'query' in query_data:
                result = self._process_single_query(query_data)
                results.extend(result)
        return results


def main():
    st.set_page_config(page_title="Auto-TKP Генератор", layout="centered", page_icon="📄")
    st.title("📄 Auto-TKP Генератор")
    st.caption("Автоматизация создания ТКП с использованием локальной LLM и RAG-поиска")

    # --- ИНИЦИАЛИЗАЦИЯ СОСТОЯНИЯ (РЕШЕНИЕ ПРОБЛЕМЫ "2-ГО РАЗА") ---
    if 'products_df' not in st.session_state:
        st.session_state.products_df = None
    if 'compatibility_df' not in st.session_state:
        st.session_state.compatibility_df = None

    st.sidebar.header("⚙️ Настройки LLM")
    llm_provider = st.sidebar.radio("Выберите сервис:", ["Ollama", "LM Studio"], index=0)

    if llm_provider == "Ollama":
        llm_client = LLMClient(provider="ollama")
    else:
        llm_client = LLMClient(provider="lm_studio")

    # Создаем экземпляры классов
    data_loader = DataLoader()
    ai_search = AISearch(llm_client, data_loader)
    doc_generator = DocumentGenerator()
    json_processor = JSONProcessor(ai_search, doc_generator)

    st.sidebar.header("📊 Статус системы")
    provider_name = "Ollama" if llm_provider == "Ollama" else "LM Studio"
    if llm_client.is_available():
        st.sidebar.success(f"✅ {provider_name} подключен")
    else:
        st.sidebar.error(f"❌ {provider_name} не отвечает")

    if st.session_state.products_df is not None:
        st.sidebar.success(f"✅ Товаров в базе: {len(st.session_state.products_df)}")
    else:
        st.sidebar.warning("⚠️ Каталог не загружен")

    st.header("1. Загрузка каталога товаров")
    st.info("💡 Загрузите `materials_final.csv`, `components.csv` и `compatibility.csv` вместе.")

    uploaded_files = st.file_uploader(
        "Выберите CSV или PDF файлы",
        type=['csv', 'pdf'],
        accept_multiple_files=True,
        key="catalog_uploader"
    )

    # Обрабатываем файлы и СОХРАНЯЕМ их в session_state
    if uploaded_files:
        # Создаем временный загрузчик для парсинга
        temp_loader = DataLoader()
        if temp_loader.load_products_from_files(uploaded_files):
            st.session_state.products_df = temp_loader.products_df
            st.session_state.compatibility_df = temp_loader.compatibility_df
            st.success(f"✅ Успешно загружено уникальных товаров: {len(st.session_state.products_df)}")
        else:
            st.error("Не удалось извлечь данные. Проверьте формат файлов.")

    st.markdown("---")
    tab1, tab2 = st.tabs(["📝 Ручной ввод запроса", "📄 Обработка JSON запросов"])

    with tab1:
        st.header("Ручной ввод запроса")
        customer_query = st.text_area(
            "Описание необходимых товаров:",
            placeholder="Например: Нужен перфорированный лоток 400 мм, крышка к нему и хомут для трубы 100 мм",
            height=80,
            key="manual_query"
        )

        order_number = st.text_input("Номер заказа", value=f"TKP-{datetime.now().strftime('%Y%m%d-%H%M')}",
                                     key="manual_order")

        if st.button("Сгенерировать ТКП", type="primary", use_container_width=True, key="manual_button"):
            # ПРОВЕРКА ИЗ SESSION_STATE (теперь она надежна на 100%)
            if st.session_state.products_df is None:
                st.error("⚠️ Сначала загрузите CSV или PDF файлы с товарами (Шаг 1)")
                st.stop()

            if not customer_query.strip():
                st.error("⚠️ Введите текст запроса")
                st.stop()

            # Передаем сохраненные данные в активный загрузчик
            data_loader.products_df = st.session_state.products_df
            data_loader.compatibility_df = st.session_state.compatibility_df

            st.warning("⏳ Первый запрос может занять 15-30 секунд (загрузка модели в память). Пожалуйста, подождите.")

            with st.spinner("🤖 AI анализирует запрос, применяет правила совместимости и подбирает товары..."):
                matches = ai_search.find_products(customer_query)

            if not matches:
                st.warning(
                    "Не найдено подходящих товаров. Попробуйте уточнить запрос или проверить наличие слов из каталога.")
                st.stop()

            st.success(f"✅ Найдено товаров: {len(matches)}")

            with st.spinner("📄 Формирование документа Word..."):
                order_data = {'order_id': order_number}
                doc, total_cost = doc_generator.generate_tcp_document(order_data, matches)

                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

            st.info(f"💰 Ориентировочная стоимость: **{total_cost} руб.**")

            st.download_button(
                label="⬇️ Скачать ТКП (.docx)",
                data=doc_buffer,
                file_name=f"{order_number}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )

    with tab2:
        st.header("Пакетная обработка JSON запросов")
        st.info("💡 Формат JSON: `[{'id': '1', 'query': 'текст'}, ...]` или `{'queries': [...]}`")

        uploaded_json_files = st.file_uploader(
            "Выберите JSON файлы с запросами",
            type=['json'],
            accept_multiple_files=True,
            key="json_uploader"
        )

        if uploaded_json_files:
            if st.session_state.products_df is None:
                st.error("⚠️ Сначала загрузите каталог товаров")
                st.stop()

            if st.button("Обработать все JSON запросы", type="primary", use_container_width=True, key="json_button"):
                data_loader.products_df = st.session_state.products_df
                data_loader.compatibility_df = st.session_state.compatibility_df

                all_results = []
                for json_file in uploaded_json_files:
                    with st.spinner(f"Обработка {json_file.name}..."):
                        results = json_processor.process_json_file(json_file)
                        all_results.extend(results)

                if all_results:
                    st.success(f"✅ Успешно обработано запросов: {len(all_results)}")
                    for result in all_results:
                        with st.expander(f"📋 Заказ {result['query_id']}: {result['query_text'][:50]}..."):
                            st.write(f"Найдено позиций: **{len(result['matches'])}**")
                            st.write(f"Общая стоимость: **{result['total_cost']} руб.**")
                            st.download_button(
                                label=f"⬇️ Скачать ТКП для {result['query_id']}",
                                data=result['document'],
                                file_name=result['filename'],
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_{result['query_id']}"
                            )
                else:
                    st.error("Не удалось обработать ни одного запроса.")


if __name__ == "__main__":
    main()